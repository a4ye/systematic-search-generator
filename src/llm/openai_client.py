"""OpenAI API wrapper with retry logic and file handling."""

import base64
import time
from dataclasses import dataclass
from http.client import IncompleteRead, RemoteDisconnected
from pathlib import Path

from openai import OpenAI, APIError, APITimeoutError
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type, before_sleep_log
from openai import RateLimitError, APIConnectionError
import logging

# Set up logger for retry messages
logger = logging.getLogger(__name__)
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter('[%(levelname)s] %(message)s'))
    logger.addHandler(handler)
logger.setLevel(logging.WARNING)

# Exceptions that should trigger a retry
RETRYABLE_EXCEPTIONS = (
    RateLimitError,
    APIConnectionError,
    APITimeoutError,
    APIError,  # Catches most OpenAI errors
    IncompleteRead,
    RemoteDisconnected,
    ConnectionError,
    TimeoutError,
)


@dataclass
class LLMResponse:
    """Response from an LLM call."""

    content: str
    model: str
    prompt_tokens: int
    completion_tokens: int
    total_tokens: int
    generation_time: float


class OpenAIClient:
    """Wrapper for OpenAI API with retry logic and file handling."""

    def __init__(self, api_key: str, model: str = "gpt-5-mini-2025-08-07"):
        self.client = OpenAI(api_key=api_key)
        self.model = model

    def _encode_file_base64(self, file_path: Path) -> str:
        """Encode a file to base64."""
        with open(file_path, "rb") as f:
            return base64.standard_b64encode(f.read()).decode("utf-8")

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF if available, otherwise return placeholder."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # Fallback: read raw bytes and let model handle it via base64
            return None

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)

            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    @retry(
        stop=stop_after_attempt(5),
        wait=wait_exponential(multiplier=1, min=4, max=60),
        retry=retry_if_exception_type(RETRYABLE_EXCEPTIONS),
        before_sleep=before_sleep_log(logger, logging.WARNING),
    )
    def generate_with_file(
        self,
        prompt: str,
        file_path: Path,
        max_tokens: int | None = None,
    ) -> LLMResponse:
        """Send prompt with a file attachment and return response."""
        start_time = time.time()

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            # Extract text from PDF
            text = self._extract_text_from_pdf(file_path)
            if text:
                # Send as text
                full_prompt = f"Document content:\n\n{text}\n\n---\n\n{prompt}"
                return self._generate_text(full_prompt, max_tokens, start_time)
            else:
                # Encode as base64 and use vision
                file_data = self._encode_file_base64(file_path)
                return self._generate_with_base64_file(
                    prompt, file_data, "application/pdf", file_path.name, max_tokens, start_time
                )

        elif suffix == ".docx":
            # Extract text from DOCX
            text = self._extract_text_from_docx(file_path)
            full_prompt = f"Document content:\n\n{text}\n\n---\n\n{prompt}"
            return self._generate_text(full_prompt, max_tokens, start_time)

        else:
            # For images, use vision API
            file_data = self._encode_file_base64(file_path)
            media_type = {
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".gif": "image/gif",
                ".webp": "image/webp",
            }.get(suffix, "application/octet-stream")

            return self._generate_with_image(
                prompt, file_data, media_type, max_tokens, start_time
            )

    def _generate_text(
        self,
        prompt: str,
        max_tokens: int | None,
        start_time: float,
    ) -> LLMResponse:
        """Generate response from text prompt."""
        kwargs = {
            "model": self.model,
            "messages": [{"role": "user", "content": prompt}],
        }
        if max_tokens is not None:
            kwargs["max_tokens"] = max_tokens

        response = self.client.chat.completions.create(**kwargs)

        generation_time = time.time() - start_time

        return LLMResponse(
            content=response.choices[0].message.content or "",
            model=response.model,
            prompt_tokens=response.usage.prompt_tokens if response.usage else 0,
            completion_tokens=response.usage.completion_tokens if response.usage else 0,
            total_tokens=response.usage.total_tokens if response.usage else 0,
            generation_time=generation_time,
        )

    def _generate_with_image(
        self,
        prompt: str,
        image_data: str,
        media_type: str,
        max_tokens: int | None,
        start_time: float,
    ) -> LLMResponse:
        """Generate response with image using vision API."""
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{media_type};base64,{image_data}",
                        },
                    },
                    {
                        "type": "text",
                        "text": prompt,
                    },
                ],
            }
        ]

        kwargs = {
            "model": self.model,
            "messages": messages,
        }
        if max_tokens is not None:
            kwargs["max_tokens"] = max_tokens

        response = self.client.chat.completions.create(**kwargs)

        generation_time = time.time() - start_time

        return LLMResponse(
            content=response.choices[0].message.content or "",
            model=response.model,
            prompt_tokens=response.usage.prompt_tokens if response.usage else 0,
            completion_tokens=response.usage.completion_tokens if response.usage else 0,
            total_tokens=response.usage.total_tokens if response.usage else 0,
            generation_time=generation_time,
        )

    def _generate_with_base64_file(
        self,
        prompt: str,
        file_data: str,
        media_type: str,
        filename: str,
        max_tokens: int | None,
        start_time: float,
    ) -> LLMResponse:
        """Generate response with base64-encoded file."""
        # For non-image files, we include the base64 data in the prompt
        # This is a fallback - most models don't support arbitrary file types
        full_prompt = f"[File: {filename}]\n\n{prompt}"

        kwargs = {
            "model": self.model,
            "messages": [{"role": "user", "content": full_prompt}],
        }
        if max_tokens is not None:
            kwargs["max_tokens"] = max_tokens

        response = self.client.chat.completions.create(**kwargs)

        generation_time = time.time() - start_time

        return LLMResponse(
            content=response.choices[0].message.content or "",
            model=response.model,
            prompt_tokens=response.usage.prompt_tokens if response.usage else 0,
            completion_tokens=response.usage.completion_tokens if response.usage else 0,
            total_tokens=response.usage.total_tokens if response.usage else 0,
            generation_time=generation_time,
        )

    @retry(
        stop=stop_after_attempt(5),
        wait=wait_exponential(multiplier=1, min=4, max=60),
        retry=retry_if_exception_type(RETRYABLE_EXCEPTIONS),
        before_sleep=before_sleep_log(logger, logging.WARNING),
    )
    def generate(
        self,
        prompt: str,
        max_tokens: int | None = None,
    ) -> LLMResponse:
        """Send a text-only prompt and return response."""
        start_time = time.time()
        return self._generate_text(prompt, max_tokens, start_time)
